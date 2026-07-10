"""
convert_to_word.py - 将result.json转换为专业剧本格式的Word文档

用法: python convert_to_word.py [result.json路径]

格式规则:
  - 场景标记(X-Y 地点 时间/内外) → Heading3
  - 对白行加粗角色名
  - 动作描写(△开头)用灰色斜体
  - 集头信息居中加粗
  - 解析 **加粗** markdown标记
  - 清理特殊符号
"""

import json
import os
import re
import sys
from pathlib import Path
from docx.shared import Pt, Inches, RGBColor, Emu


def load_result(result_path: str) -> dict:
    with open(result_path, "r", encoding="utf-8") as f:
        return json.load(f)


def get_step_content(data: dict, expert_id: str) -> str:
    """获取指定专家的内容（兼容outputs和step_outputs两种结构）"""
    outputs = data.get("outputs", {}) or data.get("step_outputs", {})
    output = outputs.get(expert_id, {})
    sd = output.get("structured_data", {})
    if isinstance(sd, dict) and sd.get("raw"):
        return sd["raw"]
    if output.get("content"):
        return output["content"]
    return ""


def _add_rich_text(paragraph, text: str, base_size=11):
    """
    解析inline markdown并添加到段落:
    - **text** → 加粗run
    - *text* → 斜体run
    - 清理特殊符号
    """
    # 清理特殊符号
    text = text.replace('✅', '·')
    text = text.replace('❌', '·')
    text = text.replace('⚠️', '·')
    text = text.replace('✨', '')
    text = text.replace('🔥', '')
    
    # 解析 **bold** 和 *italic*
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(base_size)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(base_size)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(base_size)


def create_word_doc(data: dict, output_path: str):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[错误] 需要安装python-docx: pip install python-docx")
        sys.exit(1)

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(11)

    context = data.get("context", {})
    project_name = context.get("project_name", "") or context.get("story_direction", "")[:30]
    wf_id = data.get("workflow_id", "")

    # === 封面 ===
    title = doc.add_heading(project_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("精品短剧 · 完整剧本")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("云匠 · 精品短剧创作引擎 自动生成\n工作流: " + wf_id)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_page_break()

    # === 目录 ===
    doc.add_heading("目 录", level=1)
    for item in ["一、基础信息", "二、人物小传", "三、分集大纲", "四、分场剧本", "五、对白语料库", "六、视觉方案"]:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # === 一、基础信息 ===
    doc.add_heading("一、基础信息", level=1)
    content = get_step_content(data, "§8")
    _add_content_block(doc, content) if content else doc.add_paragraph("（暂无内容）")
    doc.add_page_break()

    # === 二、人物小传 ===
    doc.add_heading("二、人物小传", level=1)
    content = get_step_content(data, "§1")
    _add_content_block(doc, content) if content else doc.add_paragraph("（暂无内容）")
    doc.add_page_break()

    # === 三、分集大纲 ===
    doc.add_heading("三、分集大纲", level=1)
    content = get_step_content(data, "§3")
    _add_content_block(doc, content) if content else doc.add_paragraph("（暂无内容）")
    doc.add_page_break()

    # === 四、分场剧本（核心）===
    doc.add_heading("四、分场剧本", level=1)
    content = get_step_content(data, "§5")
    if content:
        _add_script_block(doc, content)
    else:
        doc.add_paragraph("（暂无内容）")
    doc.add_page_break()

    # === 五、对白语料库 ===
    doc.add_heading("五、对白语料库", level=1)
    content = get_step_content(data, "§4")
    _add_content_block(doc, content) if content else doc.add_paragraph("（暂无内容）")
    doc.add_page_break()

    # === 六、视觉方案 ===
    doc.add_heading("六、视觉方案", level=1)
    content = get_step_content(data, "§13")
    _add_content_block(doc, content) if content else doc.add_paragraph("（暂无内容）")

    doc.save(output_path)
    print(f"[convert_to_word] Word文档已保存: {output_path}")


def _parse_md_table(lines: list) -> tuple:
    """解析markdown表格行，返回(headers, rows)，如果不是表格返回None"""
    if len(lines) < 3:
        return None
    
    # 第一行必须是表头
    if not lines[0].strip().startswith("|"):
        return None
    # 第二行必须是分隔线 (|---|---|)
    sep = lines[1].strip()
    if not re.match(r'^\|[\s\-:]+\|', sep):
        return None
    
    headers = [c.strip() for c in lines[0].strip().strip("|").split("|")]
    rows = []
    for line in lines[2:]:
        s = line.strip()
        if not s.startswith("|"):
            break
        cells = [c.strip() for c in s.strip("|").split("|")]
        rows.append(cells)
    
    return headers, rows, 2 + len(rows)  # 返回解析过的行数


def _add_word_table(doc, headers: list, rows: list):
    """添加真正的Word表格"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    
    num_cols = len(headers)
    num_rows = 1 + len(rows)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        # 清理markdown格式
        h_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', h)
        h_clean = re.sub(r'\*([^*]+)\*', r'\1', h_clean)
        run = p.add_run(h_clean)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "微软雅黑"
        # 表头背景色
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '4472C4',
            qn('w:val'): 'clear'
        })
        shading.append(shading_elm)
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 数据行
    for i, row in enumerate(rows):
        for j in range(min(len(row), num_cols)):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            text = row[j]
            # 清理markdown格式
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\*([^*]+)\*', r'\1', text)
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.name = "微软雅黑"
    
    doc.add_paragraph("")  # 表格后空一行


def _add_content_block(doc, content: str):
    """添加普通内容块（支持Markdown解析+表格）"""
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if not s:
            i += 1
            continue
        
        # 检测并处理表格
        if s.startswith("|"):
            # 收集连续的表格行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            # 解析并渲染表格
            parsed = _parse_md_table(table_lines)
            if parsed:
                headers, rows, _ = parsed
                _add_word_table(doc, headers, rows)
            else:
                # 不是有效表格，当普通文本
                for tl in table_lines:
                    p = doc.add_paragraph()
                    _add_rich_text(p, tl.strip())
            continue
        
        # Markdown标题
        if s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        # 列表项
        elif s.startswith("- ") or s.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            _add_rich_text(p, s[2:])
        # 分隔线
        elif s == "---" or s == "***" or (all(c == '-' for c in s) and len(s) > 3):
            pass  # 跳过分隔线
        # 普通文本（解析inline markdown）
        else:
            p = doc.add_paragraph()
            _add_rich_text(p, s)
        
        i += 1


def _add_script_block(doc, content: str):
    """
    分场剧本格式化输出（核心逻辑）
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        if not s:
            i += 1
            continue

        # 集头信息（═══包围）- 过滤掉═══线，保留内容
        if "\u2550" in s:
            i += 1
            while i < len(lines) and "\u2550" not in lines[i].strip():
                hl = lines[i].strip()
                if hl and "\u2550" not in hl:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _add_rich_text(p, hl, base_size=13)
                    for run in p.runs:
                        run.bold = True
                i += 1
            i += 1
            continue

        # 分隔线
        if s == "---" or s == "***":
            i += 1
            continue

        # 场景标记
        if re.match(r'^\d+-\d+\s+\S+.*[/][内外]', s):
            doc.add_heading(s, level=3)
            i += 1
            continue

        # 人物行
        if s.startswith("人物：") or s.startswith("人物:"):
            p = doc.add_paragraph()
            run = p.add_run(s)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)
            i += 1
            continue

        # △动作/环境描写
        if s.startswith("\u25b3"):
            p = doc.add_paragraph()
            run = p.add_run(s)
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.font.size = Pt(10)
            i += 1
            continue

        # 内心独白 (角色名os/OS：)
        m = re.match(r'^(.+?)(?:os|OS)[：:]\s*(.+)$', s)
        if m:
            p = doc.add_paragraph()
            rn = p.add_run(m.group(1).strip() + "os：")
            rn.bold = True
            rn.italic = True
            rn.font.size = Pt(10)
            _add_rich_text(p, m.group(2).strip(), base_size=10)
            for run in p.runs[1:]:
                run.italic = True
            i += 1
            continue

        # 对白行（角色名：台词）
        m = re.match(r'^(.+?)[：:]\s*(.+)$', s)
        if m:
            name = m.group(1).strip()
            text = m.group(2).strip()
            
            # 集头元信息（加粗居中）
            meta_names = ["核心事件", "情感基调", "时长", "故事方向", "一句话前提"]
            if name in meta_names:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rn = p.add_run(name + "：")
                rn.bold = True
                rn.font.size = Pt(11)
                _add_rich_text(p, text, base_size=11)
                for r in p.runs:
                    r.bold = True
                i += 1
                continue
            
            # 集标题（如：第1集《请蓝》）
            if re.match(r'^第\d+集', name):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rn = p.add_run(s)
                rn.bold = True
                rn.font.size = Pt(14)
                i += 1
                continue
            
            # 普通对白
            p = doc.add_paragraph()
            rn = p.add_run(name + "：")
            rn.bold = True
            rn.font.size = Pt(11)
            _add_rich_text(p, text, base_size=11)
            i += 1
            continue

        # 普通文本
        if s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        else:
            p = doc.add_paragraph()
            _add_rich_text(p, s)

        i += 1


def main():
    if len(sys.argv) < 2:
        candidates = ["result.json", os.path.join("workspace", "result.json")]
        result_path = None
        for c in candidates:
            if os.path.exists(c):
                result_path = c
                break
        if not result_path:
            print("用法: python convert_to_word.py [result.json路径]")
            sys.exit(1)
    else:
        result_path = sys.argv[1]

    print(f"[convert_to_word] 加载: {result_path}")
    data = load_result(result_path)

    output_dir = os.path.dirname(result_path) or "."
    wf_id = data.get("workflow_id", "output")
    output_path = os.path.join(output_dir, wf_id + "_完整剧本.docx")

    create_word_doc(data, output_path)
    print(f"[convert_to_word] 完成! 输出: {output_path}")


if __name__ == "__main__":
    main()
